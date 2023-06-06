from docx import Document
from documents.document import Document as Base
import json


class DOCXDocument(Base):

    def get_text(self):
        doc = Document(self.path)
        full_text = [para.text for para in doc.paragraphs]
        return '\n'.join(full_text)

    def set_up_file(self):
        doc = Document(self.path)
        meta = doc.core_properties
        meta.author = self.user_id
        meta.comments = json.dumps(
            {
                'control_sum': self.get_control_sum(self.get_text()),
                'signers': []
            }
        )
        doc.save(self.path)

    def sign(self):
        if self.validate():
            doc = Document(self.path)
            core = doc.core_properties
            meta = json.loads(core.comments)
            meta['signers'].append(self.user_id)
            core.comments = json.dumps(meta)
            doc.save(self.path)

    def validate(self):
        try:
            doc = Document(self.path)
            cs = json.loads(doc.core_properties.comments)['control_sum']
        except (KeyError, json.decoder.JSONDecodeError):
            raise ValueError('Document has never been initialized')
        if cs == self.get_control_sum(self.get_text()):
            return True
        raise ValueError('Document has been changed. Control sum doesnt suit the content')

    def who_signed(self):
        if self.validate():
            return json.loads(Document(self.path).core_properties.comments)['signers']

    def is_signed_by(self):
        return self.user_id in self.who_signed()
